import logging
from pathlib import Path

from docx import Document as DocxDocument
from apps.parsing.choices import ExtractionMethod
from apps.parsing.parsers.base import (
    DocumentParseResult,
    ParsedPageResult,
    ParsedSectionResult,
    ParsedTableResult,
)
from apps.parsing.services.quality import aggregate_quality, score_page_text
from apps.parsing.services.section_detection import (
    _extract_heading_title,
    _is_heading_line,
    build_structured_text,
)

logger = logging.getLogger(__name__)


def _heading_level(paragraph) -> int | None:
    style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
    if style_name.startswith("heading"):
        try:
            return int(style_name.replace("heading", "").strip() or "1")
        except ValueError:
            return 1
    return None


def parse_docx(file_path: Path) -> DocumentParseResult:
    doc = DocxDocument(str(file_path))
    paragraphs_text: list[str] = []
    blocks: list[tuple[str, str]] = []  # (type, text) type: heading|para

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        level = _heading_level(para)
        if level is not None:
            blocks.append(("heading", text))
        elif _is_heading_line(text):
            blocks.append(("heading", _extract_heading_title(text)))
        else:
            blocks.append(("para", text))
        paragraphs_text.append(text)

    # Tables in reading order after body (POC: append at end)
    tables: list[ParsedTableResult] = []
    for table_index, table in enumerate(doc.tables):
        rows_raw: list[list[str]] = []
        for row in table.rows:
            rows_raw.append([cell.text.strip() for cell in row.cells])
        if not rows_raw:
            continue
        headers = rows_raw[0]
        data_rows = rows_raw[1:] if len(rows_raw) > 1 else []
        tables.append(
            ParsedTableResult(
                page_number=1,
                headers=headers,
                rows=data_rows,
                raw=rows_raw,
            )
        )
        table_text = _table_to_text(headers, data_rows)
        blocks.append(("table", table_text))
        paragraphs_text.append(table_text)

    full_text = "\n".join(paragraphs_text)
    quality = score_page_text(full_text)

    pages = [
        ParsedPageResult(
            page_number=1,
            extracted_text=full_text,
            extraction_method=ExtractionMethod.DOCX_NATIVE,
            ocr_used=False,
            quality_score=quality,
            is_empty=not full_text.strip(),
        )
    ]

    sections = _sections_from_blocks(blocks)
    structured_text = build_structured_text(sections)

    metadata = {
        "file_type": "docx",
        "parser": "python-docx",
        "total_pages": 1,
        "empty_pages": 1 if not full_text.strip() else 0,
        "ocr_pages": 0,
        "tables": [
            {"page": t.page_number, "headers": t.headers, "rows": t.rows} for t in tables
        ],
        "page_quality": [
            {
                "page": 1,
                "quality_score": quality,
                "extraction_method": ExtractionMethod.DOCX_NATIVE,
                "ocr_used": False,
            }
        ],
    }

    return DocumentParseResult(
        pages=pages,
        sections=sections,
        tables=tables,
        raw_text=full_text,
        structured_text=structured_text,
        parsing_metadata=metadata,
        parsing_quality_score=quality,
        file_type="docx",
    )


def _sections_from_blocks(blocks: list[tuple[str, str]]) -> list[ParsedSectionResult]:
    sections: list[ParsedSectionResult] = []
    current_title = "Preamble"
    current_lines: list[str] = []
    order = 0

    def flush() -> None:
        nonlocal order, current_title, current_lines
        content = "\n".join(current_lines).strip()
        if content or order == 0:
            sections.append(
                ParsedSectionResult(
                    title=current_title,
                    content=content,
                    page_start=1,
                    page_end=1,
                    section_order=order,
                )
            )
            order += 1
        current_lines = []

    for kind, text in blocks:
        if kind == "heading":
            flush()
            current_title = text
        else:
            current_lines.append(text)

    flush()

    if not sections:
        return [
            ParsedSectionResult(
                title="Document",
                content="",
                page_start=1,
                page_end=1,
                section_order=0,
            )
        ]
    return sections


def _table_to_text(headers: list[str], rows: list[list[str]]) -> str:
    lines = [" | ".join(headers)]
    for row in rows:
        lines.append(" | ".join(row))
    return "\n".join(lines)
